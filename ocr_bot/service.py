try:
    from PIL import Image
except ImportError:
    import Image
# unrar lib: https://www.rarlab.com/rar/unrarsrc-5.9.2.tar.gz
from rarfile import RarFile
import io
import tempfile
import pytesseract
from docx import Document
from docx.shared import Pt

def convert(image_data: bytearray, lang='rus', mime_type=None) -> bytes:
    bits = io.BytesIO(image_data)
    if not mime_type:
        return _convert_image_to_text(bits, lang).encode('utf-8'), 'txt'
    if mime_type == 'application/vnd.rar':
        fp = tempfile.mktemp()
        with open(fp, 'wb') as temp_file:
            temp_file.write(bits.getvalue())

        doc = Document()

        rar = RarFile(fp)
        for pic in rar.infolist():
            bits = rar.open(pic.filename)
            par = doc.add_paragraph('')
            par.paragraph_format.page_break_before = True
            run = par.add_run(
                _convert_image_to_text(bits, lang)
            )
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        
        fout = io.BytesIO()
        doc.save(fout)
        return fout.getvalue(), 'docx'


def _convert_image_to_text(bits, lang):
    img = Image.open(bits)
    return pytesseract.image_to_string(img, lang=lang)